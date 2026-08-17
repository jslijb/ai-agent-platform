#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""读取需求说明书 docx 内容并保存为 txt（便于分析）"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    import os
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document

docx_path = Path("需求说明书V3.1_外发交流版_免NDA_20260727(1).docx")
output_path = Path("需求说明书V3.1_extracted.txt")

doc = Document(str(docx_path))

lines = []
# 段落
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else ""
        lines.append(f"[{style}] {text}")

# 表格
for ti, table in enumerate(doc.tables):
    lines.append(f"\n=== 表格 {ti+1} ===")
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append(" | ".join(cells))

content = "\n".join(lines)
output_path.write_text(content, encoding="utf-8")
print(f"Extracted {len(lines)} lines, {len(content)} chars to {output_path}")
print(f"\nFirst 50 lines:")
for line in lines[:50]:
    print(line)
