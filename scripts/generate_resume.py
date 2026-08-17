from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

FONT_BODY = Pt(10.5)
FONT_SMALL = Pt(9)
FONT_H1 = Pt(14)
FONT_H2 = Pt(11)
FONT_NAME = '微软雅黑'
LINE_SPACING = 1.15

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = FONT_BODY
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = LINE_SPACING

for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

def set_run(run, size=FONT_BODY, bold=False, color=None):
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run('▎' + text)
    set_run(run, size=FONT_H2, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run('─' * 70)
    run2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run2.font.size = Pt(6)

def add_para(text, bold=False, size=FONT_BODY, indent=0, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p

def add_mixed_para(parts, indent=0, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold, size in parts:
        run = p.add_run(text)
        set_run(run, size=size or FONT_BODY, bold=bold)
    return p

def add_bullet(text, indent=0.4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.2)
    run = p.add_run('• ' + text)
    set_run(run, size=FONT_BODY)
    return p

# ==================== 标题 ====================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(4)
p_title.paragraph_format.space_before = Pt(0)
run_t = p_title.add_run('个 人 简 历')
set_run(run_t, size=FONT_H1, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))

# ==================== 基本信息 ====================
add_section_title('基本信息')

info_lines = [
    [('姓　名：', True), ('________', False)],
    [('性　别：', True), ('________', False), ('　　', False), ('出生年月：', True), ('________', False)],
    [('手机号码：', True), ('________', False), ('　　', False), ('电子邮箱：', True), ('________', False)],
    [('现居城市：', True), ('________', False), ('　　', False), ('工作年限：', True), ('________', False)],
    [('求职状态：', True), ('________', False)],
]
for line_parts in info_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    for text, bold in line_parts:
        run = p.add_run(text)
        set_run(run, size=FONT_SMALL, bold=bold)

# ==================== 期望职位 ====================
add_section_title('期望职位')
add_para('全栈工程师 / Python 工程师 / Agent 工程师', size=FONT_SMALL)

# ==================== 个人优势 ====================
add_section_title('个人优势')

advantages = [
    'Claude Code + SDD/TDD：Agentic workflow · 多文件重构 · MCP集成 · 200K上下文 —— 用规范驱动拆解复杂任务，用红-绿-重构闭环保障代码质量。',
    'Agent：LangGraph、OpenClaw、多智能体协同、Skill编排、工具调用、MCP风格',
    '大数据框架：Airflow、Flink-CDC、Hive、Spark、数据调度',
    'RAG：MultiRAG、Milvus、Qdrant、FAISS、BGE、多模态检索、混合召回',
    '大模型部署+调优+测评：Qwen3-32B、LoRA微调、SFT、MMLU/C-Eval/Fin-Eva评估、benchmark设计',
    '系统部署与集成：Docker、Docker Compose、多阶段构建、镜像瘦身、容器编排',
    'DevOps/CI/CD：Prometheus、cAdvisor、日志健康检查、GitOps、自动化测试与审批流水线',
]
for adv in advantages:
    add_bullet(adv)

# ==================== 工作经历 ====================
add_section_title('工作经历')

add_mixed_para([
    ('深圳市法本信息技术有限公司', True, FONT_BODY),
    ('　　2018-10 — 2026-01　　', False, FONT_SMALL),
    ('数据架构师 / Agent开发工程师', True, FONT_BODY),
])

add_bullet('负责企业级大数据平台建设：湖仓一体（Iceberg/Hudi）+ Flink流批一体 + 多维数据库（Kylin/Doris）+ BI报表体系（Superset/FineBI）')
add_bullet('主导数据治理与指标体系搭建，通过Redis多层缓存优化查询性能')
add_bullet('主导智能BI编程Agent（DataAgent）从0到1落地，实现自然语言→SQL/代码→调度→报表生成+数据/报表治理全链路')
add_bullet('设计多智能体协作、RAG知识库（业务知识+SQL案例）、Redis缓存加速、血缘依赖检测、数据质量保障')

add_mixed_para([('工作业绩：', True, FONT_BODY)], space_after=Pt(1))
add_bullet('PB级数据支撑，实时报表延迟<5分钟，多维查询<2秒，BI性能提升40%，开发周期缩短60%')
add_bullet('Agent复杂自助取数占比85%，SQL准确率89%，相比传统开发节省人力40人天/月，下线冗余资产节约存储20%')

# ==================== 项目经历 ====================
add_section_title('项目经历')

# --- 项目一 ---
add_mixed_para([
    ('金融AI智能体平台', True, FONT_BODY),
    ('　　项目架构师　　', False, FONT_SMALL),
    ('2025-07 — 至今', False, FONT_SMALL),
    ('　　（兼职项目）', False, FONT_SMALL),
])

add_mixed_para([('项目描述：', True, FONT_BODY)], space_after=Pt(1))
add_para('构建一站式金融AI智能体平台，集成多Agent协作、RAG/GraphRAG检索、MCP工具协议等核心能力，支持智能投研、量化分析、合规审查。经历V1原型→V2生产级的完整演进，涵盖基础设施迁移、Agent架构重构、RAG质量治理、微服务拆分四大升级。', size=FONT_SMALL)

add_mixed_para([('技术栈：', True, FONT_BODY)], space_after=Pt(1))
add_para('Next.js 14 + TypeScript + Route Handlers(SSE) + NextAuth v5 + PostgreSQL(pgvector) + Neo4j + Redis + LangGraph.js + 阿里百炼 + Python FastAPI + Docker Compose', size=FONT_SMALL)

add_para('一、基础设施迁移', bold=True, size=FONT_BODY)
add_bullet('ORM：Prisma → Drizzle ORM，解决影子库superuser权限依赖、Rust引擎构建慢、pgvector兼容性不足三大痛点')
add_bullet('API层：tRPC → Route Handlers + SSE，原生支持流式推送与中间件集成')
add_bullet('数据缓存：SQLite → PostgreSQL PgCache，消除文件锁和跨服务数据不一致')

add_para('二、Agent分层编排', bold=True, size=FONT_BODY)
add_bullet('单体SimpleAgent 21工具平铺→多Agent编排（Researcher/Quant/Compliance）+ Skill技能层')
add_bullet('「Query → Skill → Tool」三层决策，13+声明式Skill固化高频任务，Prompt Token↓50%+')
add_bullet('多工具链式执行（3轮→2轮，省30% Token）+ 重复调用检测 + 数据真实性校验 + 反思循环')
add_bullet('四层分层记忆（L1原始/L2滚动摘要/L3向量检索/L4用户画像）+ 自适应Token预算')

add_para('三、RAG质量治理（Top-5准确率↑40%+）', bold=True, size=FONT_BODY)
add_bullet('清洗管线：控制字符→Markdown噪声→全半角归一化→Unicode NFC标准化')
add_bullet('智能切片：800字符 + 128重叠 + 句子边界感知 + 多级断点（丢失率36%→<5%）')
add_bullet('混合检索：BM25稀疏 + pgvector稠密 + RRF(K=60)融合；分离精排（文档Top-5 + 图谱Top-3）')
add_bullet('查询增强：HyDE假设文档改写 + 金融同义词扩展')
add_bullet('合规自动过期：研报90天/年报365天/法规永不过期，全链路引用溯源')

add_para('四、GraphRAG知识图谱', bold=True, size=FONT_BODY)
add_bullet('LLM自动抽取金融关系三元组，Neo4j多跳推理检索')
add_bullet('CDC监听增量同步Embedding/图谱/BM25索引，图谱限流防噪声')

add_para('五、金融视觉分析', bold=True, size=FONT_BODY)
add_bullet('双引擎策略：PaddleOCR主力（本地SOTA）→ Qwen-VL降级保可用')
add_bullet('3个视觉Skill：研报截图结构化、K线形态识别、财报OCR指标计算')

add_para('六、MCP工具生态', bold=True, size=FONT_BODY)
add_bullet('10个MCP标准化工具 + 21+ Agent内部工具，统一ToolRegistry注册管理')
add_bullet('内部ToolRegistry高性能调用 + 外部MCP SSE标准化暴露，双轨设计')
add_bullet('工具描述增强（when_to_use / when_not_to_use / few-shot）+ 调用校验层')

add_para('七、生产级稳定性', bold=True, size=FONT_BODY)
add_bullet('LLM降级链：开发环境基于api_keys.yaml驱动多模型切换；生产环境部署私有化大模型（72B主力+14B降级），vLLM推理引擎+多副本负载均衡+GPU健康检查自动恢复')
add_bullet('temperature=0 + seed=42确定性输出；LLM语义缓存（TTL 30min）；IP滑动窗口限流')
add_bullet('多级降级：Reranker→原始排序；图谱→跳过；Redis→内存缓存；HNSW→顺序扫描')

add_para('八、微服务架构', bold=True, size=FONT_BODY)
add_bullet('5个自研服务 + Nginx网关：主服务(Next.js) + RAG(:3001) + LLM Gateway(:3002) + Evaluation(:3003) + Data Service(:8001 Python)')
add_bullet('ServiceAdapter单体/微服务双模式，USE_MICROSERVICE=false一键回退')
add_bullet('BullMQ异步评估队列 + TraceId分布式追踪 + Prometheus + Grafana可观测')

add_para('九、评估体系', bold=True, size=FONT_BODY)
add_bullet('RAG 10指标（数值准确率/合规/幻觉/风险/时效等）+ Agent 5指标（工具选择/规划/合规/一致性/效率）')
add_bullet('适配FinEval/CFLUE/FinQA/ConvFinQA四类开源数据集 + 103条黄金测试集回归')

add_mixed_para([('项目业绩：', True, FONT_BODY)], space_after=Pt(1))
add_bullet('混合检索+重排Top-5准确率↑40%+；切片优化丢失率36%→<5%')
add_bullet('Agent重构复杂查询迭代轮次↓40%+，Prompt Token↓50%+')
add_bullet('微服务拆分实现故障隔离，评估OOM不影响主站，RAG/LLM独立扩容')
add_bullet('完整日志、熔断降级、语义缓存、分布式追踪与评估体系')

# --- 项目二 ---
add_mixed_para([
    ('DataAgent——湖仓一体智能数据开发与治理平台', True, FONT_BODY),
    ('　　项目架构师　　', False, FONT_SMALL),
    ('2025-05 — 2026-01', False, FONT_SMALL),
])

add_mixed_para([('技术栈：', True, FONT_BODY)], space_after=Pt(1))
add_para('OpenClaw / LangGraph / Flink / Paimon / Kafka / Redis / Kylin / Doris / Milvus / BGE / FastAPI / Airflow', size=FONT_SMALL)

add_para('● 智能数据开发与报表生成', bold=True, size=FONT_BODY)
add_bullet('自然语言需求 → Agent规划拆解 → 自动构建数仓逻辑层（DWD明细层、DWS汇总层）→ 生成SQL/Python代码 → 调度配置 → BI报表自动创建（调用Superset/FineBI API）')
add_bullet('Agent根据业务需求自动设计事实表、维度表、汇总指标，输出符合数仓规范的DDL/DML，确保数据模型可复用、易治理')
add_bullet('支持运营、产品、分析师以自然语言驱动数仓演进，自动完成从业务口径到物理模型的映射')

add_para('● 多智能体协作', bold=True, size=FONT_BODY)
add_bullet('部署需求分析、代码生成、测试、报表生成、部署、治理6个专用Agent，通过消息队列协同，人工仅最终审批')

add_para('● RAG知识库增强', bold=True, size=FONT_BODY)
add_bullet('存储数据字典、业务术语映射、指标口径，让Agent理解业务逻辑')
add_bullet('记录大模型生成SQL的成功/失败案例，失败案例入库自动修正Prompt；沉淀SQL标准化规范（命名、性能、安全规则）')
add_bullet('基于Milvus+BGE向量检索，实现长期记忆与上下文注入')

add_para('● Redis缓存加速Agent链路', bold=True, size=FONT_BODY)
add_bullet('缓存会话历史与短期记忆，降低LLM token消耗')
add_bullet('缓存工具调用结果（Schema查询、维度数据），减少重复外部调用')
add_bullet('存储Agent中间规划结果，支持断点续跑与异常恢复')

add_para('● 报表治理', bold=True, size=FONT_BODY)
add_bullet('根据报表访问频率、最后使用时间，自动识别长期无人访问报表，推送下线建议或自动下线')
add_bullet('由报表反推数据模型：若某模型仅被待下线报表使用，联动评估下线该模型')

add_para('● 血缘依赖检测与数据治理', bold=True, size=FONT_BODY)
add_bullet('基于OpenLineage解析SQL血缘，在建议下线模型前自动检测下游依赖，避免影响在用报表')
add_bullet('持续扫描数仓，识别冗余数据模型（相同逻辑宽表、未使用中间表），给出合并/下线建议')
add_bullet('解决开发痛点：开发人员修改表时，Agent通过血缘自动找到所有下游依赖，代替人工沟通，降低故障风险')

add_para('● 数据质量检测与BI健壮性', bold=True, size=FONT_BODY)
add_bullet('代码生成阶段自动注入DQC规则（空值、唯一性、量级波动）')
add_bullet('监控BI报表数据新鲜度、异常值、SQL错误率，异常时自动告警/回滚/通知')

add_para('● 推动生产上线', bold=True, size=FONT_BODY)
add_bullet('先小范围独立集群测试，经过多轮迭代，稳定后依旧是独立集群小团队试验，主要面向运营、产品、数据分析师、数据挖掘师。')

add_mixed_para([('项目业绩：', True, FONT_BODY)], space_after=Pt(1))
add_bullet('Agent SQL准确率89%（对比资深工程师91%）')
add_bullet('报表自动生成占新报表总量60%')
add_bullet('血缘检测避免15次模型误删故障，下游查找时间从2小时→5分钟')

# ==================== 教育经历 ====================
add_section_title('教育经历')
add_para('武昌首义学院　　全日制大专　　专业：计算机应用工程　　2004 — 2007', size=FONT_SMALL)

output_path = r'D:\Python\ai-agent-platform\docs\个人简历.docx'
doc.save(output_path)
print(f'简历已保存到: {output_path}')
